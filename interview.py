
import gc
import time
import google.generativeai as genai
import pdfplumber
import threading
import json
import re
import os
import tempfile
import sys
import subprocess
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict
from collections import defaultdict
import random
from datetime import datetime

# Imports for Word Document Generation
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# Imports for Core Functionality
import speech_recognition as sr
from gtts import gTTS
import pygame
import cv2
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk

# --- 2. CONFIGURATION ---
from dotenv import load_dotenv
load_dotenv()  # Load environment variables from .env file

API_KEY = os.getenv('GEMINI_API_KEY_2')  # Get API key from environment variable
if not API_KEY:
    print("ðŸ›‘ ERROR: GEMINI_API_KEY_2 not found in environment variables")

# Configure the API
model = None
try:
    if not API_KEY or "YOUR_VALID_API_KEY_HERE" in API_KEY:
        print("ðŸ›‘ ERROR: Please replace 'YOUR_VALID_API_KEY_HERE' with your actual Gemini API key.")
    else:
        genai.configure(api_key=API_KEY)
        model = genai.GenerativeModel('gemini-2.5-flash')
        print("âœ… Gemini API Key configured and model initialized.")
except Exception as e:
    print(f"ðŸ›‘ ERROR: Could not configure Gemini API. Please check your key. Error: {e}")
    print("âš   Continuing without API access - some features will not work")

# --- 3. DATA STRUCTURES ---
@dataclass
class Question:
    type: str
    question: str
    category: str = "general"  # More specific categorization

@dataclass
class Evaluation:
    score: int
    feedback: str
    suggestions: str
    corrected_answer: str
    keywords_matched: List[str] = None

@dataclass
class InterviewPlan:
    branch: str
    skills_summary: str
    projects_summary: str
    question_bank: List[Question]
    resume_keywords: List[str] = None

@dataclass
class InterviewResult:
    question: Question
    user_answer: str
    evaluation: Evaluation
    face_detected: bool = True

# Initialize pygame for audio playback with better error handling
try:
    pygame.mixer.init(frequency=22050, size=-16, channels=2, buffer=512)
    print("âœ… Pygame mixer initialized successfully.")
except pygame.error as e:
    print(f"ðŸ›‘ Pygame mixer could not be initialized. Audio playback will be disabled. Error: {e}")

# Initialize speech recognition with better configuration
recognizer = sr.Recognizer()
recognizer.energy_threshold = 300
recognizer.pause_threshold = 1.5
recognizer.dynamic_energy_threshold = True

microphone = None
try:
    microphone = sr.Microphone()
    print("Calibrating microphone for ambient noise...")
    with microphone as source:
        recognizer.adjust_for_ambient_noise(source, duration=3)
    print("âœ… Microphone calibration complete.")
except Exception as e:
    print(f"ðŸ›‘ Microphone not found or could not be initialized. Voice input will not work. Error: {e}")

# --- 4. HELPER FUNCTIONS ---
def get_info_from_gemini(prompt: str, is_json_response: bool = True, max_retries: int = 3) -> Optional[Any]:
    """Get and parse responses from the Gemini API with retry logic."""
    if not model:
        print("âš   Gemini model not available, using fallback.")
        if "interview plan" in prompt: return get_fallback_interview_plan()
        return None

    for attempt in range(max_retries):
        try:
            response = model.generate_content(prompt)
            text = response.text.strip()
            
            if not is_json_response:
                return text

            # Try to extract JSON from response
            json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', text, re.DOTALL)
            if json_match:
                json_text = json_match.group(1)
            elif text.startswith('{') and text.endswith('}'):
                json_text = text
            else:
                # Attempt to find the JSON block if formatting is off
                start_index = text.find('{')
                end_index = text.rfind('}') + 1
                if start_index != -1 and end_index != 0:
                    json_text = text[start_index:end_index]
                else:
                    raise ValueError("No valid JSON object found in the response.")
            
            return json.loads(json_text)
        except Exception as e:
            print(f"ðŸ›‘ AI Response Error (Attempt {attempt+1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
            else:
                return None

def get_fallback_interview_plan():
    """Fallback interview plan when Gemini API is not available."""
    return {
        "branch": "Computer Science (Fallback)",
        "skills_summary": "Programming, Problem Solving",
        "projects_summary": "Software development projects",
        "question_bank": [
            {"type": "resume", "question": "Tell me about your most significant project."},
            {"type": "core", "question": "Explain a core concept from Data Structures."},
            {"type": "core", "question": "Write a simple SQL query to find all users from a 'users' table in 'New York'."},
            {"type": "HR", "question": "Where do you see yourself in 5 years?"}
        ]
    }

def speak_question(text: str) -> None:
    """Converts question text to speech and plays it automatically with better error handling."""
    temp_file_path = None
    try:
        print(f"ðŸ”Š Speaking question: {text}")
        tts = gTTS(text=text, lang='en', slow=False)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
            temp_file_path = fp.name
            tts.save(temp_file_path)
        
        # Wait a moment for file to be fully written
        time.sleep(0.5)
        
        pygame.mixer.music.load(temp_file_path)
        pygame.mixer.music.play()
        
        # Wait for playback to finish
        while pygame.mixer.music.get_busy():
            pygame.time.wait(100)
            
    except Exception as e:
        print(f"ðŸ›‘ Could not play audio: {e}")
        # Fallback: print the question loudly
        print("ðŸ”Š SPEAKING QUESTION (FALLBACK):", text.upper())
    finally:
        # Clean up temp file
        if temp_file_path and os.path.exists(temp_file_path):
            for attempt in range(3):
                try:
                    os.unlink(temp_file_path)
                    break
                except PermissionError:
                    if attempt < 2:
                        time.sleep(0.5)
                    else:
                        print(f"âš  Warning: Could not delete temp audio file: {temp_file_path}")

def record_audio(timeout: int = 15, phrase_time_limit: int = 60) -> Optional[str]:
    """Record audio from microphone and convert to text with improved reliability."""
    if not microphone:
        print("ðŸŽ¤ Microphone is not available. Skipping audio recording.")
        return "Microphone not available, no answer recorded."
    
    try:
        print(f"\nðŸŽ¤ Recording... Speak now! (Timeout: {timeout}s, Max phrase: {phrase_time_limit}s)")
        print("   Listening for your answer...")
        
        with microphone as source:
            # Recalibrate for current noise level
            recognizer.adjust_for_ambient_noise(source, duration=1)
            audio_data = recognizer.listen(source, timeout=timeout, phrase_time_limit=phrase_time_limit)
        
        print("   Processing your speech...")
        text = recognizer.recognize_google(audio_data)
        print(f"âœ… Recognized: \"{text}\"")
        return text
    
    except sr.WaitTimeoutError:
        print("   â° No speech detected within the timeout period.")
        return "No speech detected."
    except sr.UnknownValueError:
        print("   ðŸ¤” Could not understand the audio. Please try again.")
        return "Audio not understood."
    except sr.RequestError as e:
        print(f"   ðŸ”Œ Error with speech recognition service: {e}")
        return "Speech recognition service error."
    except Exception as e:
        print(f"   ðŸ›‘ Unexpected error recording audio: {e}")
        return f"Recording error: {str(e)}"

def setup_camera():
    """Set up camera for face detection and verification."""
    try:
        cap = cv2.VideoCapture(0)
        if not cap.isOpened():
            print("âŒ Camera not accessible.")
            return None
        
        # Set camera resolution for better performance
        cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
        cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
        cap.set(cv2.CAP_PROP_FPS, 30)
        
        print("âœ… Camera access successful!")
        return cap
    except Exception as e:
        print(f"Error accessing camera: {e}")
        return None

def detect_face(cap):
    """Detect if a face is present in the camera frame."""
    try:
        if cap is None or not cap.isOpened():
            return False
            
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        ret, frame = cap.read()
        if not ret: 
            return False
        
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        faces = face_cascade.detectMultiScale(gray, 1.1, 4)
        return len(faces) > 0
    except Exception as e:
        print(f"Face detection error: {e}")
        return False

# --- 5. CORE LOGIC ---
def parse_resume(pdf_path: str) -> Optional[str]:
    """Extracts text from a PDF using pdfplumber."""
    print("ðŸ“„ Parsing resume...")
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            print("âš  Warning: Could not extract meaningful text from the PDF.")
            return None
        
        print("âœ… Resume parsed successfully.")
        return text
    except Exception as e:
        print(f"ðŸ›‘ ERROR: Failed to parse PDF file. {e}")
        return None

def extract_keywords_from_resume(resume_text: str, max_keywords: int = 15) -> List[str]:
    """Extract important keywords from resume using AI."""
    prompt = f"""
    Analyze the following resume text and extract the {max_keywords} most important technical skills, 
    technologies, and domain-specific keywords. Return ONLY a JSON array of strings.
    
    Resume Text: --- {resume_text[:3000]} ---
    """
    
    keywords = get_info_from_gemini(prompt)
    if keywords and isinstance(keywords, list):
        return keywords[:max_keywords]
    
    # Fallback to simple keyword extraction if AI fails
    technical_terms = {
        'python', 'java', 'javascript', 'sql', 'html', 'css', 'react', 'angular', 'vue', 
        'node', 'express', 'django', 'flask', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
        'machine learning', 'ai', 'data science', 'analytics', 'backend', 'frontend', 'fullstack',
        'devops', 'ci/cd', 'agile', 'scrum', 'rest', 'api', 'graphql', 'microservices'
    }
    
    found_keywords = []
    for term in technical_terms:
        if term in resume_text.lower():
            found_keywords.append(term)
    
    return found_keywords[:max_keywords]

def get_branch_from_resume(resume_text: str) -> str:
    """Uses AI to determine the candidate's academic/engineering branch."""
    prompt = f"""
    Analyze the following resume text and identify the candidate's primary academic or engineering branch.
    Respond with ONLY the name of the branch (e.g., "Computer Science", "Mechanical Engineering", "Business Administration").
    
    Resume Text: --- {resume_text[:2000]} ---
    """
    branch = get_info_from_gemini(prompt, is_json_response=False)
    return branch or "General"

def generate_interview_plan(resume_text: str, previous_questions: List[str] = None) -> Optional[InterviewPlan]:
    """Uses AI to create a complete, structured interview plan tailored to the candidate's branch."""

    print("\nðŸ§  Determining candidate's branch from resume...")
    branch_name = get_branch_from_resume(resume_text) or "Unknown"
    print(f"âœ… Detected Branch: {branch_name}")

    # Extract keywords for question variety
    resume_keywords = extract_keywords_from_resume(resume_text) or []
    print(f"ðŸ“‹ Extracted Keywords: {', '.join(resume_keywords) if resume_keywords else 'None'}")

    # Branch-wise core subject prompts
    branch_core_map = {
    "computer": "Generate 1 question on Operating Systems, 1 question on Computer Networks,1 question on Object-Oriented Programming System, 1 questions on Data Structures & Algorithms, and 1 SQL question.",
    "it": "Generate 1 question on Operating Systems, 1 question on Computer Networks,1 question on Object-Oriented Programming System, 1 questions on Data Structures & Algorithms, and 1 SQL question.",
    "csd": "Generate 1 question on Operating Systems, 1 question on Computer Networks,1 question on Object-Oriented Programming System, 1 questions on Data Structures & Algorithms, and 1 SQL question.",
    "aids": "Generate 1 question on Operating Systems, 1 question on Computer Networks,1 question on Object-Oriented Programming System, 1 questions on Data Structures & Algorithms, and 1 SQL question.",
    "civil": "Generate 3 questions from Civil Engineering core subjects (Structural, Surveying, Geotechnical, Concrete).",
    "mechanical": "Generate 3 questions from Mechanical core subjects (Thermodynamics, SOM, Machine Design).",
    "electrical": "Generate 3 questions from Electrical core subjects (Machines, Power Systems, Circuits).",
    "entc": "Generate 3 questions from ENTC core subjects (Signals, Microprocessors, Communication Systems).",
    "chemical": "Generate 3 questions from Chemical core subjects (Heat Transfer, Fluid Mechanics, Reaction Engineering).",
    "robotics": "Generate 3 questions from Robotics core subjects (Kinematics, Sensors, Control Systems)."
}

    # Pick the key (first word lowercase) to match map
    branch_key = branch_name.lower().split()[0]
    avoid_questions = ""
    if previous_questions:
        avoid_questions = f" Avoid these questions that were asked before: {', '.join(previous_questions[:5])}."

    core_questions_prompt = branch_core_map.get(branch_key, f"""
        Generate 3 in-depth questions relevant to {branch_name} branch core subjects.
    """) + avoid_questions

    # Final AI prompt
    final_prompt = f"""
    Analyze the following resume text for a candidate from the '{branch_name}' branch and create a complete interview plan.
    Focus on these key areas: {', '.join(resume_keywords) if resume_keywords else 'general skills'}.

    Your response MUST be a single, valid JSON object.
    The JSON object should have these exact keys: "branch", "skills_summary", "projects_summary", "question_bank".

    For the "question_bank", create a list of 7 diverse question objects with "type" and "question" keys:
    1. Generate 2 questions based directly on the projects or skills mentioned in the resume (type: "resume").
    2. {core_questions_prompt} (type: "core").
    3. Generate 2 standard HR questions to evaluate soft skills (type: "HR").

    Resume Text: --- {resume_text[:3500]} ---
    """

    print("ðŸ§  Generating tailored interview plan with AI...")
    result = get_info_from_gemini(final_prompt)

    if result and isinstance(result, dict) and "branch" in result and "question_bank" in result:
        print("âœ… Tailored interview plan generated successfully.")
        questions = [
            Question(type=q.get("type", "N/A").lower(), question=q.get("question", ""))
            for q in result.get("question_bank", [])
        ]
        return InterviewPlan(
            branch=result.get("branch", branch_name),
            skills_summary=result.get("skills_summary", "N/A"),
            projects_summary=result.get("projects_summary", "N/A"),
            question_bank=questions,
            resume_keywords=resume_keywords
        )

    print("ðŸ›‘ Failed to generate interview plan from AI response.")
    return None

def get_user_answer(question_type: str, camera=None) -> tuple[str, bool]:
    """Get user answer through voice with camera monitoring."""
    stop_monitoring = threading.Event()
    detection_results = []

    def monitor_camera():
        nonlocal detection_results
        while not stop_monitoring.is_set():
            if camera and detect_face(camera):
                detection_results.append(True)
            else:
                detection_results.append(False)
            time.sleep(0.5)  # Check more frequently for better accuracy

    monitor_thread = threading.Thread(target=monitor_camera, daemon=True)
    if camera: 
        monitor_thread.start()

    answer = record_audio()

    stop_monitoring.set()
    if camera: 
        monitor_thread.join(timeout=2.0)

    # Calculate face detection percentage
    if detection_results:
        detection_rate = sum(detection_results) / len(detection_results)
        final_face_detected = detection_rate >= 0.6  # 60% detection threshold
    else:
        final_face_detected = True  # If no camera, assume face detected
    
    return answer or "No answer provided", final_face_detected

def evaluate_answer(question: Question, answer: str, resume_keywords: List[str] = None) -> Evaluation:
    """Evaluate the user's answer using AI."""
    keywords_context = ""
    if resume_keywords:
        keywords_context = f"Pay special attention to these keywords from the resume: {', '.join(resume_keywords)}."
    
    prompt = f"""
    You are an expert interviewer. Evaluate this interview answer on a scale of 1-10.
    The response must be a valid JSON object.
    
    Question: "{question.question}"
    Candidate's Answer: "{answer}"
    {keywords_context}

    Provide a JSON response with these keys:
    - "score": An integer from 1 to 10.
    - "feedback": A concise, constructive paragraph on the answer's strengths and weaknesses.
    - "suggestions": A short bulleted list (as a single string with '\\n- ') on how to improve.
    - "corrected_answer": An example of an ideal, well-structured answer to the question.
    - "keywords_matched": A list of keywords from the resume that were effectively addressed in the answer (if any).
    """
    
    result = get_info_from_gemini(prompt)
    if result and all(k in result for k in ["score", "feedback", "suggestions", "corrected_answer"]):
        # Ensure keywords_matched is always a list
        if "keywords_matched" not in result:
            result["keywords_matched"] = []
        return Evaluation(**result)
    
    # Fallback evaluation
    return Evaluation(
        score=5,
        feedback="Could not get AI evaluation. This is a fallback.",
        suggestions="- Try to be more specific.\n- Provide more detailed examples.",
        corrected_answer="An ideal answer would provide more detail and specific examples.",
        keywords_matched=[]
    )

def generate_report_summary(results: List[InterviewResult]) -> Dict[str, Any]:
    """Uses AI to generate a high-level summary of the interview."""
    # Format questions and answers for the AI
    qa_pairs = []
    for i, result in enumerate(results):
        qa_pairs.append(f"Q{i+1}: {result.question.question}\nA: {result.user_answer}\nScore: {result.evaluation.score}/10")
    
    qa_text = "\n\n".join(qa_pairs)
    
    prompt = f"""
    Analyze these interview question-answer pairs and provide a comprehensive summary.
    Return a JSON object with these keys:
    - "overall_summary": A paragraph summarizing the candidate's overall performance.
    - "key_strengths": A list of 3-5 key strengths demonstrated.
    - "areas_for_improvement": A list of 3-5 areas needing improvement.
    - "recommendation": One of: "Strongly Recommend", "Recommend", "Neutral", "Not Recommend".
    
    Interview Q&A:
    {qa_text}
    """
    
    result = get_info_from_gemini(prompt)
    if result:
        return result
    
    # Fallback summary
    total_score = sum(r.evaluation.score for r in results)
    avg_score = total_score / len(results) if results else 0
    
    return {
        "overall_summary": f"The candidate demonstrated basic knowledge with an average score of {avg_score:.1f}/10.",
        "key_strengths": ["Clear communication", "Relevant experience"],
        "areas_for_improvement": ["Could provide more specific examples", "Should demonstrate deeper technical knowledge"],
        "recommendation": "Neutral"
    }

def add_formatted_paragraph(document, text, bold=False, italic=False, color=None, font_size=12):
    """Add a formatted paragraph to the document."""
    p = document.add_paragraph()
    run = p.add_run(text)
    
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_size:
        run.font.size = Pt(font_size)
    
    return p

def add_colored_table_row(table, cells, colors=None):
    """Add a row to a table with optional cell colors."""
    row_cells = table.add_row().cells
    for i, cell_text in enumerate(cells):
        row_cells[i].text = str(cell_text)
        
        # Apply color if specified
        if colors and i < len(colors) and colors[i]:
            shading_elm = parse_xml(f'<w:shd {qn("w:fill")}="{colors[i]}"/>')
            row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

def generate_word_report(interview_plan: InterviewPlan, interview_results: List[InterviewResult], camera_verified: bool):
    """Generate a professional Word (.docx) report with detailed evaluation."""
    document = docx.Document()
    
    # Set document properties
    document.core_properties.title = "AI Interview Performance Report"
    document.core_properties.author = "AI Interview System"
    document.core_properties.comments = f"Interview for {interview_plan.branch} candidate"
    
    # Set page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header with styling
    heading = document.add_heading('Interview Performance Report', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        run.font.size = Pt(16)
        run.bold = True

    # Subtitle with date
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle_run.font.size = Pt(12)
    subtitle_run.italic = True
    document.add_paragraph()

    # Performance Summary
    document.add_heading('Executive Summary', level=1)
    summary_data = generate_report_summary(interview_results)
    
    # Overall summary
    add_formatted_paragraph(document, "Overall Assessment:", bold=True, font_size=12)
    add_formatted_paragraph(document, summary_data['overall_summary'], font_size=11)
    document.add_paragraph()

    # Key strengths
    add_formatted_paragraph(document, "Key Strengths:", bold=True, font_size=12)
    for strength in summary_data['key_strengths']:
        add_formatted_paragraph(document, f"â€¢ {strength}", font_size=11)
    document.add_paragraph()

    # Areas for improvement
    add_formatted_paragraph(document, "Areas for Improvement:", bold=True, font_size=12)
    for improvement in summary_data['areas_for_improvement']:
        add_formatted_paragraph(document, f"â€¢ {improvement}", font_size=11)
    document.add_paragraph()

    # Recommendation
    rec_color = {
        "Strongly Recommend": (0, 128, 0),      # Green
        "Recommend": (0, 102, 204),             # Blue
        "Neutral": (255, 153, 0),               # Orange
        "Not Recommend": (204, 0, 0)            # Red
    }.get(summary_data['recommendation'], (0, 0, 0))
    
    add_formatted_paragraph(document, "Final Recommendation:", bold=True, font_size=12)
    add_formatted_paragraph(document, summary_data['recommendation'], color=rec_color, font_size=12, bold=True)
    document.add_paragraph()

    # Interview Overview
    document.add_heading('Interview Details', level=1)
    total_score = sum(result.evaluation.score for result in interview_results)
    avg_score = total_score / len(interview_results) if interview_results else 0
    
    # Create a table for overview
    table = document.add_table(rows=0, cols=2)
    table.style = 'LightShading-Accent1'
    
    # Add overview data
    overview_data = [
        ("Candidate Branch", interview_plan.branch),
        ("Skills Focus", interview_plan.skills_summary),
        ("Projects Focus", interview_plan.projects_summary),
        ("Average Score", f"{avg_score:.1f} / 10"),
        ("Camera Verification", "Passed" if camera_verified else "Failed"),
        ("Total Questions", len(interview_results))
    ]
    
    for key, value in overview_data:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
        row_cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()

    # Performance by Category
    document.add_heading('Performance Analysis', level=1)
    
    # Category scores table
    scores_by_type = defaultdict(lambda: {'total': 0, 'count': 0})
    for res in interview_results:
        cat = res.question.type.capitalize()
        scores_by_type[cat]['total'] += res.evaluation.score
        scores_by_type[cat]['count'] += 1
    
    cat_table = document.add_table(rows=1, cols=3)
    cat_table.style = 'LightShading-Accent1'
    hdr_cells = cat_table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Questions'
    hdr_cells[2].text = 'Average Score'
    
    for cell in hdr_cells: 
        cell.paragraphs[0].runs[0].bold = True
    
    for category, data in scores_by_type.items():
        if data['count'] > 0:
            avg = data['total'] / data['count']
            row_cells = cat_table.add_row().cells
            row_cells[0].text = category
            row_cells[1].text = str(data['count'])
            row_cells[2].text = f"{avg:.1f} / 10"
    
    document.add_paragraph()
    
    # Keyword analysis if available
    if interview_plan.resume_keywords:
        document.add_heading('Keyword Coverage Analysis', level=2)
        
        # Count keyword mentions in evaluations
        keyword_mentions = {}
        for keyword in interview_plan.resume_keywords:
            mention_count = sum(1 for result in interview_results 
                              if keyword.lower() in result.evaluation.keywords_matched)
            keyword_mentions[keyword] = mention_count
        
        # Create keyword table
        kw_table = document.add_table(rows=0, cols=2)
        kw_table.style = 'LightList-Accent1'
        
        kw_header = kw_table.add_row().cells
        kw_header[0].text = 'Resume Keyword'
        kw_header[1].text = 'Times Addressed'
        for cell in kw_header: 
            cell.paragraphs[0].runs[0].bold = True
        
        for keyword, count in sorted(keyword_mentions.items(), key=lambda x: x[1], reverse=True):
            row_cells = kw_table.add_row().cells
            row_cells[0].text = keyword
            row_cells[1].text = str(count)
    
    document.add_page_break()

    # Detailed Question Analysis
    document.add_heading('Detailed Question Analysis', level=1)
    
    for i, result in enumerate(interview_results):
        # Question header
        q_header = document.add_heading(f"Question {i+1}: {result.question.type.capitalize()} Question", level=2)
        for run in q_header.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Question text
        add_formatted_paragraph(document, "Question:", bold=True)
        add_formatted_paragraph(document, f'"{result.question.question}"', italic=True)
        
        # Candidate's answer
        add_formatted_paragraph(document, "Candidate's Answer:", bold=True)
        add_formatted_paragraph(document, f'"{result.user_answer}"')
        
        # Score with color coding
        score_color = (
            (0, 128, 0) if result.evaluation.score >= 8 else  # Green for high scores
            (0, 102, 204) if result.evaluation.score >= 6 else  # Blue for medium scores
            (204, 0, 0)  # Red for low scores
        )
        add_formatted_paragraph(document, "Score:", bold=True)
        add_formatted_paragraph(document, f"{result.evaluation.score}/10", color=score_color, bold=True)
        
        # Feedback
        add_formatted_paragraph(document, "Feedback:", bold=True)
        add_formatted_paragraph(document, result.evaluation.feedback)
        
        # Suggestions
        add_formatted_paragraph(document, "Suggestions for Improvement:", bold=True)
        add_formatted_paragraph(document, result.evaluation.suggestions)
        
        # Ideal answer
        add_formatted_paragraph(document, "Example of an Ideal Answer:", bold=True)
        ideal_para = document.add_paragraph()
        ideal_para.paragraph_format.left_indent = Inches(0.25)
        ideal_run = ideal_para.add_run(result.evaluation.corrected_answer)
        ideal_run.italic = True
        ideal_run.font.color.rgb = RGBColor(0, 102, 0)
        
        # Keyword matches
        if result.evaluation.keywords_matched:
            add_formatted_paragraph(document, "Keywords Addressed:", bold=True)
            add_formatted_paragraph(document, ", ".join(result.evaluation.keywords_matched))
        
        # Face detection status
        status_color = (0, 128, 0) if result.face_detected else (204, 0, 0)
        status_text = "Face Detected" if result.face_detected else "Face Not Detected"
        add_formatted_paragraph(document, "Verification:", bold=True)
        add_formatted_paragraph(document, status_text, color=status_color)
        
        document.add_paragraph()  # Add spacing between questions

    # Save the Document
    report_filename = f"AI_Interview_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    document.save(report_filename)
    print(f"âœ… Word report successfully generated: {report_filename}")
    return report_filename, avg_score

# --- 6. GUI INTERFACE ---
class InterviewApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Interview System")
        self.root.geometry("900x700")
        self.root.configure(bg="#f5f7f9")
        self.resume_path = None
        self.resume_text = None
        self.interview_plan = None
        self.camera = None
        self.interview_results = []
        self.face_detection_results = []
        self.previous_questions = self.load_previous_questions()
        self.camera_running = False
        self.setup_ui()

    def load_previous_questions(self):
        """Load questions from previous interviews to ensure variety."""
        try:
            if os.path.exists("interview_history.json"):
                with open("interview_history.json", "r") as f:
                    data = json.load(f)
                    return data.get("previous_questions", [])
        except:
            pass
        return []

    def save_interview_history(self):
        """Save current questions to history for future variety."""
        try:
            current_questions = [q.question for q in self.interview_plan.question_bank] if self.interview_plan else []
            
            # Load existing history
            if os.path.exists("interview_history.json"):
                with open("interview_history.json", "r") as f:
                    data = json.load(f)
            else:
                data = {"previous_questions": []}
            
            # Add new questions and keep only recent ones (avoid storing too many)
            data["previous_questions"].extend(current_questions)
            data["previous_questions"] = data["previous_questions"][-50:]  # Keep last 50 questions
            
            with open("interview_history.json", "w") as f:
                json.dump(data, f)
        except Exception as e:
            print(f"Error saving interview history: {e}")

    def setup_ui(self):
        # Main container
        main_container = tk.Frame(self.root, bg="#f5f7f9")
        main_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Header
        header_frame = tk.Frame(main_container, bg="#f5f7f9")
        header_frame.pack(pady=(0, 30))
        
        tk.Label(header_frame, text="AI Interview System", font=("Arial", 28, "bold"), 
                bg="#f5f7f9", fg="#2c3e50").pack(pady=(0, 10))
        
        tk.Label(header_frame, text="Upload your resume to start a tailored interview", 
                font=("Arial", 14), bg="#f5f7f9", fg="#7f8c8d").pack()
        
        # Upload section
        upload_frame = tk.Frame(main_container, bg="#f5f7f9")
        upload_frame.pack(pady=(0, 30))
        
        self.upload_btn = tk.Button(upload_frame, text="ðŸ“„ Upload Resume (PDF)", 
                                   command=self.upload_resume, font=("Arial", 12), 
                                   bg="#3498db", fg="white", padx=25, pady=12,
                                   activebackground="#2980b9")
        self.upload_btn.pack()
        
        self.resume_status = tk.Label(upload_frame, text="No resume uploaded", 
                                     font=("Arial", 10), bg="#f5f7f9", fg="#95a5a6")
        self.resume_status.pack(pady=(15, 0))
        
        # Start button
        self.start_btn = tk.Button(main_container, text="ðŸŽ¤ Start Interview", 
                                  command=self.start_interview, font=("Arial", 16, "bold"), 
                                  bg="#27ae60", fg="white", padx=40, pady=15, 
                                  state=tk.DISABLED, activebackground="#219a52")
        self.start_btn.pack(pady=(20, 0))
        
        # Progress bar
        self.progress = ttk.Progressbar(main_container, mode='indeterminate', length=400)
        self.progress.pack(fill=tk.X, pady=(30, 0))
        
        # Status label
        self.status_label = tk.Label(main_container, text="", font=("Arial", 11), 
                                    bg="#f5f7f9", fg="#34495e", wraplength=600)
        self.status_label.pack(pady=(15, 0))
        
        # Footer
        footer_frame = tk.Frame(main_container, bg="#f5f7f9")
        footer_frame.pack(side=tk.BOTTOM, pady=(20, 0))
        
        tk.Label(footer_frame, text="Uses AI-powered evaluation â€¢ Camera verification â€¢ Professional reporting", 
                font=("Arial", 9), bg="#f5f7f9", fg="#bdc3c7").pack()

    def upload_resume(self):
        file_path = filedialog.askopenfilename(
            title="Select your resume PDF", 
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if not file_path: 
            return

        self.resume_path = file_path
        self.resume_status.config(text=f"Uploaded: {os.path.basename(file_path)}", fg="#16a085")
        self.start_btn.config(state=tk.DISABLED)
        self.progress.start(10)
        self.status_label.config(text="Analyzing resume and generating interview questions...", fg="#34495e")
        self.root.update()
        
        # Process in background thread
        threading.Thread(target=self.parse_and_analyze, daemon=True).start()

    def parse_and_analyze(self):
        self.resume_text = parse_resume(self.resume_path)
        if self.resume_text:
            self.interview_plan = generate_interview_plan(self.resume_text, self.previous_questions)
            self.root.after(0, self.on_resume_analyzed)
        else:
            self.root.after(0, self.on_resume_error)

    def on_resume_analyzed(self):
        self.progress.stop()
        if self.interview_plan:
            self.status_label.config(text="Resume analyzed successfully! Interview plan generated.", fg="#27ae60")
            self.start_btn.config(state=tk.NORMAL)
            
            # Show brief overview
            overview = f"Branch: {self.interview_plan.branch} | Questions: {len(self.interview_plan.question_bank)}"
            self.resume_status.config(text=f"{self.resume_status.cget('text')} | {overview}")
            
            messagebox.showinfo("Ready", 
                f"Interview plan generated for a {self.interview_plan.branch} role with {len(self.interview_plan.question_bank)} questions.\n\n"
                f"Press 'Start Interview' when you are ready. The interview will use your microphone and camera.")
        else:
            self.status_label.config(text="Could not generate interview plan from resume.", fg="#e74c3c")
            messagebox.showerror("Error", 
                "Failed to create an interview plan. The AI may have been unable to understand the resume content.\n"
                "Please ensure your PDF contains selectable text and try again.")

    def on_resume_error(self):
        self.progress.stop()
        self.status_label.config(text="Error analyzing resume. Please try again with a different PDF.", fg="#e74c3c")
        messagebox.showerror("Error", 
            "Failed to analyze the resume. Please make sure:\n"
            "1. It's a valid PDF file\n"
            "2. It contains selectable text (not just images)\n"
            "3. The file is not corrupted or password protected")

    def start_interview(self):
        if not self.interview_plan:
            messagebox.showerror("Error", "Please upload and analyze a resume first.")
            return
        
        confirm = messagebox.askyesno("Confirm", 
            "The interview will:\n"
            "â€¢ Use your microphone for voice answers\n"
            "â€¢ Use your camera for verification\n"
            "â€¢ Ask tailored questions based on your resume\n\n"
            "Make sure you're in a quiet environment with good lighting.\n\n"
            "Ready to begin?")
            
        if not confirm:
            return
        
        # Clear the window and show interview screen
        for widget in self.root.winfo_children(): 
            widget.destroy()
        
        self.show_interview_in_progress()
        threading.Thread(target=self.run_interview, daemon=True).start()

    def show_interview_in_progress(self):
        self.root.configure(bg="#2c3e50")
        
        # Header
        header_frame = tk.Frame(self.root, bg="#2c3e50")
        header_frame.pack(fill=tk.X, pady=(20, 10))
        
        tk.Label(header_frame, text="AI Interview In Progress", font=("Arial", 24, "bold"), 
                bg="#2c3e50", fg="#ecf0f1").pack()
        
        # Main content frame
        content_frame = tk.Frame(self.root, bg="#2c3e50")
        content_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Left side - Question and Answer
        left_frame = tk.Frame(content_frame, bg="#34495e", relief=tk.RAISED, bd=1)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # Question display
        question_frame = tk.Frame(left_frame, bg="#34495e")
        question_frame.pack(fill=tk.X, padx=20, pady=20)
        
        self.question_number_label = tk.Label(question_frame, text="", 
                                             font=("Arial", 16, "bold"), 
                                             bg="#34495e", fg="#3498db")
        self.question_number_label.pack(pady=(0, 10))
        
        self.question_label = tk.Label(question_frame, text="Preparing first question...", 
                                      font=("Arial", 18), bg="#34495e", fg="#ecf0f1", 
                                      wraplength=500, justify=tk.CENTER)
        self.question_label.pack(pady=10)
        
        # Recording status
        self.recording_status = tk.Label(question_frame, text="", 
                                        font=("Arial", 14), 
                                        bg="#34495e", fg="#f39c12")
        self.recording_status.pack(pady=(10, 0))
        
        # Answer display area
        answer_frame = tk.Frame(left_frame, bg="#2c3e50")
        answer_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        tk.Label(answer_frame, text="Your Answer (Voice to Text):", font=("Arial", 14, "bold"), 
                bg="#2c3e50", fg="#ecf0f1").pack(anchor=tk.W, pady=(0, 10))
        
        self.answer_text = tk.Text(answer_frame, height=8, wrap=tk.WORD, 
                                  font=("Arial", 12), bg="#1a252f", fg="#ecf0f1",
                                  insertbackground="white")
        self.answer_text.pack(fill=tk.BOTH, expand=True)
        self.answer_text.config(state=tk.DISABLED)
        
        # Right side - Camera feed
        right_frame = tk.Frame(content_frame, bg="#34495e", relief=tk.RAISED, bd=1)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=(10, 0))
        
        tk.Label(right_frame, text="Camera Feed - Face Verification", font=("Arial", 14, "bold"), 
                bg="#34495e", fg="#ecf0f1").pack(pady=20)
        
        self.camera_label = tk.Label(right_frame, bg="#1a252f", width=40, height=25)
        self.camera_label.pack(padx=20, pady=20)
        
        # Status bar at bottom
        status_frame = tk.Frame(self.root, bg="#2c3e50")
        status_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.interview_status_label = tk.Label(status_frame, text="Initializing...", font=("Arial", 12), 
                                              bg="#2c3e50", fg="#bdc3c7")
        self.interview_status_label.pack()

    def update_ui_threadsafe(self, widget, config):
        self.root.after(0, lambda: widget.config(**config))

    def update_camera_preview(self):
        """Update camera preview in GUI continuously."""
        if self.camera_running and self.camera and self.camera.isOpened():
            try:
                ret, frame = self.camera.read()
                if ret:
                    # Detect face and draw rectangle
                    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
                    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                    faces = face_cascade.detectMultiScale(gray, 1.1, 4)
                    
                    for (x, y, w, h) in faces:
                        cv2.rectangle(frame, (x, y), (x+w, y+h), (0, 255, 0), 2)
                        cv2.putText(frame, "Face Detected", (x, y-10), 
                                  cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 2)
                    
                    # Resize and convert for display
                    frame = cv2.resize(frame, (320, 240))
                    frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    img = Image.fromarray(frame)
                    imgtk = ImageTk.PhotoImage(image=img)
                    
                    self.camera_label.imgtk = imgtk
                    self.camera_label.configure(image=imgtk)
                
                # Continue updating
                if self.camera_running:
                    self.root.after(50, self.update_camera_preview)
            except Exception as e:
                print(f"Camera preview error: {e}")
                if self.camera_running:
                    self.root.after(50, self.update_camera_preview)

    def run_interview(self):
        """Main interview execution loop with improved voice functionality."""
        # Start camera
        self.camera = setup_camera()
        self.camera_running = True
        
        if self.camera:
            self.root.after(0, self.update_camera_preview)
            self.update_ui_threadsafe(self.interview_status_label, {
                "text": "Camera started - Face verification active"
            })
        else:
            self.update_ui_threadsafe(self.interview_status_label, {
                "text": "Camera not available - Continuing without face verification"
            })
        
        total_questions = len(self.interview_plan.question_bank)
        
        for i, question in enumerate(self.interview_plan.question_bank):
            q_num = i + 1
            
            # Update UI with current question
            self.update_ui_threadsafe(self.question_number_label, {
                "text": f"Question {q_num} of {total_questions}"
            })
            self.update_ui_threadsafe(self.question_label, {
                "text": question.question
            })
            self.update_ui_threadsafe(self.recording_status, {
                "text": "ðŸ”Š Speaking question..."
            })
            self.update_ui_threadsafe(self.interview_status_label, {
                "text": f"Question {q_num} of {total_questions} | Speaking question..."
            })
            
            # Clear previous answer
            self.update_ui_threadsafe(self.answer_text, {"state": tk.NORMAL})
            self.answer_text.delete("1.0", tk.END)
            self.update_ui_threadsafe(self.answer_text, {"state": tk.DISABLED})
            
            # Speak the question (voice) - Wait for it to complete
            speak_question(question.question)
            
            # Update status for recording
            self.update_ui_threadsafe(self.recording_status, {
                "text": "ðŸŽ¤ Listening for your answer... Speak now!"
            })
            self.update_ui_threadsafe(self.interview_status_label, {
                "text": f"Question {q_num} of {total_questions} | Listening for your answer..."
            })
            
            # Get user answer with camera monitoring (voice input)
            answer, face_detected = get_user_answer(question.type, self.camera)
            
            # Display the recognized answer in the text box
            if answer:
                self.update_ui_threadsafe(self.answer_text, {"state": tk.NORMAL})
                self.answer_text.delete("1.0", tk.END)
                self.answer_text.insert("1.0", answer)
                self.update_ui_threadsafe(self.answer_text, {"state": tk.DISABLED})
            
            self.update_ui_threadsafe(self.recording_status, {
                "text": "ðŸ¤” Evaluating your answer..."
            })
            self.update_ui_threadsafe(self.interview_status_label, {
                "text": f"Question {q_num} of {total_questions} | Evaluating your answer..."
            })
            
            # Evaluate the answer
            evaluation = evaluate_answer(question, answer, self.interview_plan.resume_keywords)
            self.interview_results.append(InterviewResult(question, answer, evaluation, face_detected))
            
            # Show immediate feedback
            self.update_ui_threadsafe(self.recording_status, {
                "text": f"âœ… Score: {evaluation.score}/10"
            })
            self.update_ui_threadsafe(self.interview_status_label, {
                "text": f"Question {q_num} of {total_questions} | Score: {evaluation.score}/10"
            })
            
            # Brief pause between questions
            time.sleep(2)

        # Clean up camera
        self.camera_running = False
        if self.camera:
            self.camera.release()
            cv2.destroyAllWindows()
            self.update_ui_threadsafe(self.camera_label, {"image": ""})

        # Save interview history for future variety
        self.save_interview_history()

        # Generate final report
        self.update_ui_threadsafe(self.interview_status_label, {
            "text": "Interview complete! Generating your detailed report..."
        })
        
        overall_face_detected = all(result.face_detected for result in self.interview_results)
        report_filename, avg_score = generate_word_report(
            self.interview_plan, self.interview_results, overall_face_detected
        )
        
        self.root.after(0, self.show_interview_complete, report_filename, avg_score)

    def show_interview_complete(self, report_filename, avg_score):
        """Show the interview completion screen with results."""
        for widget in self.root.winfo_children(): 
            widget.destroy()
        
        self.root.configure(bg="#ecf0f1")
        
        # Completion message
        tk.Label(self.root, text="Interview Complete!", font=("Arial", 28, "bold"), 
                bg="#ecf0f1", fg="#2c3e50").pack(pady=(60, 20))
        
        # Score display
        score_color = (
            "#27ae60" if avg_score >= 8 else
            "#3498db" if avg_score >= 6 else
            "#e74c3c"
        )
        
        tk.Label(self.root, text=f"Overall Score: {avg_score:.1f}/10", 
                font=("Arial", 20), bg="#ecf0f1", fg=score_color).pack(pady=(0, 30))
        
        # Report location
        if report_filename:
            report_path = os.path.abspath(report_filename)
            tk.Label(self.root, text="A detailed report has been generated:", 
                    font=("Arial", 12), bg="#ecf0f1", fg="#7f8c8d").pack(pady=(0, 5))
            
            tk.Label(self.root, text=report_path, font=("Arial", 10), 
                    bg="#ecf0f1", fg="#3498db", wraplength=600).pack(pady=(0, 30))
        else:
            tk.Label(self.root, text="Report generation failed. Please check the console for errors.", 
                    font=("Arial", 12), bg="#ecf0f1", fg="#e74c3c").pack(pady=(0, 30))
        
        # Buttons
        button_frame = tk.Frame(self.root, bg="#ecf0f1")
        button_frame.pack(pady=(0, 40))
        
        if report_filename and os.path.exists(report_filename):
            tk.Button(button_frame, text="Open Report", 
                     command=lambda: self.open_report(report_filename),
                     font=("Arial", 12), bg="#3498db", fg="white", padx=20, pady=10).pack(side=tk.LEFT, padx=10)
        
        tk.Button(button_frame, text="New Interview", 
                 command=self.restart_app,
                 font=("Arial", 12), bg="#2ecc71", fg="white", padx=20, pady=10).pack(side=tk.LEFT, padx=10)
        
        tk.Button(button_frame, text="Exit", 
                 command=self.root.destroy,
                 font=("Arial", 12), bg="#e74c3c", fg="white", padx=20, pady=10).pack(side=tk.LEFT, padx=10)

    def open_report(self, report_filename):
        """Open the generated report file."""
        try:
            if os.name == 'nt':  # Windows
                os.startfile(report_filename)
            elif os.name == 'posix':  # macOS, Linux
                if sys.platform == "darwin":
                    subprocess.run(["open", report_filename])
                else:
                    subprocess.run(["xdg-open", report_filename])
            else:
                messagebox.showinfo("Report", f"Report saved to:\n{report_filename}")
        except Exception as e:
            messagebox.showinfo("Report", f"Report saved to:\n{report_filename}\n\nError opening file: {e}")

    def restart_app(self):
        """Restart the application for a new interview."""
        self.camera_running = False
        if self.camera:
            self.camera.release()
        self.root.destroy()
        main()

# --- 7. MAIN EXECUTION ---
def main():
    """Main entry point for the application."""
    try:
        root = tk.Tk()
        app = InterviewApp(root)
        root.mainloop()
    except Exception as e:
        print(f"ðŸ›‘ Fatal error: {e}")
        messagebox.showerror("Error", f"The application encountered a fatal error:\n{e}")

if __name__ == "__main__":
    main()