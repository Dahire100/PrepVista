import re
import os
import warnings
import json
from docx import Document
from docx.shared import Inches, RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
import pdfplumber
import google.generativeai as genai

# Suppress warnings
warnings.filterwarnings('ignore')

print('Professional Resume Analyzer loaded successfully.')

# ---------------------- Gemini API Setup ----------------------
from dotenv import load_dotenv
load_dotenv()  # Load environment variables from .env file

GEMINI_API_KEY = os.getenv('GEMINI_API_KEY_1')  # Get API key from environment variable
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY_1 not found in environment variables")

genai.configure(api_key=GEMINI_API_KEY)

# ---------------------- Professional Formatting Functions ----------------------
def set_cell_background(cell, fill_color):
    """Set cell background color with better compatibility"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_style(doc, text, level=1, color=None, top_spacing=6, bottom_spacing=6):
    """Add a styled heading with spacing options"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)

    # Add spacing
    heading.paragraph_format.space_before = Pt(top_spacing)
    heading.paragraph_format.space_after = Pt(bottom_spacing)

    return heading

def create_score_card(doc, title, score, details, color="4472C4"):
    """Create a professional score card with details"""
    # Title row
    title_table = doc.add_table(rows=1, cols=1)
    title_table.autofit = False
    title_table.columns[0].width = Inches(6.0)
    title_cell = title_table.cell(0, 0)
    title_cell.text = title
    title_cell.paragraphs[0].runs[0].font.bold = True
    title_cell.paragraphs[0].runs[0].font.size = Pt(12)
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_background(title_cell, color)
    title_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Score and details table
    score_table = doc.add_table(rows=1, cols=2)
    score_table.autofit = False
    score_table.columns[0].width = Inches(1.5)
    score_table.columns[1].width = Inches(4.5)

    # Score cell
    score_cell = score_table.cell(0, 0)
    score_cell.text = f"{score}%"
    score_cell.paragraphs[0].runs[0].font.bold = True
    score_cell.paragraphs[0].runs[0].font.size = Pt(20)

    # Determine score color
    if score >= 80:
        score_color = RGBColor(46, 125, 50)  # Green
    elif score >= 60:
        score_color = RGBColor(239, 108, 0)  # Orange
    else:
        score_color = RGBColor(198, 40, 40)  # Red

    score_cell.paragraphs[0].runs[0].font.color.rgb = score_color
    score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Details cell
    details_cell = score_table.cell(0, 1)
    for detail in details:
        p = details_cell.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(detail)
        run.font.size = Pt(10)

    doc.add_paragraph()  # Add spacing

# ---------------------- Core Functions ----------------------
def extract_text_from_pdf(pdf_path):
    """Extract text from PDF with error handling"""
    text = ''
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n\n'
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""
    return text.strip()

def analyze_resume_with_gemini(resume_text, jd_text):
    """Analyze resume and job description using Gemini API"""
    try:
        model = genai.GenerativeModel("gemini-2.0-flash-exp")
        prompt = f"""
        You are a professional resume analysis tool. Your task is to analyze a resume and a job description and provide a detailed report.

        **Instructions:**
        1.  Analyze the provided RESUME and JOB DESCRIPTION texts.
        2.  Generate a comprehensive JSON object with the following structure:
            - **candidate_info**: A dictionary with 'name', 'email', 'phone', 'linkedin', 'location', and 'summary' extracted from the resume. If a field is not found, use "Not Found".
            - **ats_score**: A score from 0-100 based on general ATS best practices.
            - **ats_details**: A list of short, bullet-point sentences explaining how the score was calculated.
            - **match_score**: A score from 0-100 indicating how well the resume matches the job description.
            - **match_details**: A list of key matching points or mismatches.
            - **improvement_tips**: A list of 5-8 specific, actionable tips for improving the resume for this job.
            - **keywords_found**: A list of keywords from the job description found in the resume.
            - **keywords_missing**: A list of important keywords from the job description that were not found in the resume.
            - **strengths**: A list of 3-5 key strengths of the candidate for this role.
            - **weaknesses**: A list of 3-5 potential weaknesses or gaps for this role.
            - **skills_assessment**: A dictionary with skills as keys and scores (0-100) as values for at least 6 key skills mentioned in the job description.
            - **recommendation**: One of: "Strongly Recommend", "Recommend", "Consider with Reservations", "Not Recommended"
        3.  Ensure the JSON is well-formed and valid. Do not include any text before or after the JSON object.

        **RESUME TEXT:**
        {resume_text}

        **JOB DESCRIPTION TEXT:**
        {jd_text}
        """

        response = model.generate_content(prompt, stream=False)

        # Extract and clean the JSON from the response
        json_text = response.text.replace("```json", "").replace("```", "").strip()
        analysis_results = json.loads(json_text)

        return analysis_results
    except Exception as e:
        print(f"Error calling Gemini API: {e}")
        return None

def generate_report(analysis_results, output_path='Resume_Analysis_Report.docx'):
    """Generate comprehensive professional report in DOCX format"""

    if not analysis_results:
        print("Invalid analysis results. Cannot generate report.")
        return None

    # Ensure output_path ends with .docx
    if not output_path.endswith('.docx'):
        output_path = output_path.replace('.pdf', '.docx')

    doc = Document()

    # Set document properties
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ---------------------- TITLE PAGE ----------------------
    # Main title
    title = doc.add_heading("RESUME ANALYSIS REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(23, 54, 93)

    # Subtitle
    subtitle = doc.add_paragraph("Professional Candidate Assessment")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Date
    date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(11)
    date_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # ---------------------- CANDIDATE INFORMATION ----------------------
    add_heading_with_style(doc, "CANDIDATE PROFILE", 1, (23, 54, 93), 12, 12)

    info = analysis_results.get('candidate_info', {})

    # Contact information table
    contact_table = doc.add_table(rows=5, cols=2)
    contact_table.style = 'Table Grid'
    contact_table.autofit = False
    contact_table.columns[0].width = Inches(1.8)
    contact_table.columns[1].width = Inches(4.2)

    contact_data = [
        ("Full Name", info.get('name', 'Not Provided')),
        ("Email Address", info.get('email', 'Not Provided')),
        ("Phone Number", info.get('phone', 'Not Provided')),
        ("LinkedIn Profile", info.get('linkedin', 'Not Provided')),
        ("Location", info.get('location', 'Not Provided'))
    ]

    for i, (label, value) in enumerate(contact_data):
        cells = contact_table.rows[i].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(11)

        cells[1].text = value
        cells[1].paragraphs[0].runs[0].font.size = Pt(11)

        # Color code based on presence of information
        if value == 'Not Provided' or value == 'Not Found':
            set_cell_background(cells[1], "FFEBEE")
        else:
            set_cell_background(cells[1], "E8F5E9")

    doc.add_paragraph()

    # ---------------------- EXECUTIVE SUMMARY ----------------------
    add_heading_with_style(doc, "EXECUTIVE SUMMARY", 1, (23, 54, 93), 24, 12)

    ats_score = analysis_results.get('ats_score', 0)
    match_score = analysis_results.get('match_score', 0)

    # Create summary table
    summary_table = doc.add_table(rows=2, cols=2)
    summary_table.autofit = False
    summary_table.columns[0].width = Inches(2.8)
    summary_table.columns[1].width = Inches(3.2)

    # ATS Score
    ats_label = summary_table.cell(0, 0)
    ats_label.text = "ATS Compliance"
    ats_label.paragraphs[0].runs[0].font.bold = True

    ats_value = summary_table.cell(0, 1)
    ats_value.text = f"{ats_score}%"
    if ats_score >= 80:
        ats_value.paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 125, 50)
    elif ats_score >= 60:
        ats_value.paragraphs[0].runs[0].font.color.rgb = RGBColor(239, 108, 0)
    else:
        ats_value.paragraphs[0].runs[0].font.color.rgb = RGBColor(198, 40, 40)

    # Match Score
    match_label = summary_table.cell(1, 0)
    match_label.text = "Job Match"
    match_label.paragraphs[0].runs[0].font.bold = True

    match_value = summary_table.cell(1, 1)
    match_value.text = f"{match_score}%"
    if match_score >= 80:
        match_value.paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 125, 50)
    elif match_score >= 60:
        match_value.paragraphs[0].runs[0].font.color.rgb = RGBColor(239, 108, 0)
    else:
        match_value.paragraphs[0].runs[0].font.color.rgb = RGBColor(198, 40, 40)

    # Overall recommendation
    doc.add_paragraph()
    if ats_score >= 80 and match_score >= 80:
        recommendation = "STRONG CANDIDATE - Resume is well-optimized and shows excellent alignment with position requirements."
        color = (46, 125, 50)
    elif ats_score >= 60 and match_score >= 60:
        recommendation = "MODERATE CANDIDATE - Good foundation with some areas for improvement to enhance competitiveness."
        color = (239, 108, 0)
    else:
        recommendation = "NEEDS IMPROVEMENT - Significant revisions recommended to meet position requirements."
        color = (198, 40, 40)

    rec_para = doc.add_paragraph()
    rec_para.paragraph_format.space_before = Pt(12)
    rec_para.paragraph_format.space_after = Pt(12)
    run = rec_para.add_run(recommendation)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(12)

    # ---------------------- DETAILED ANALYSIS ----------------------
    add_heading_with_style(doc, "DETAILED ANALYSIS", 1, (23, 54, 93), 24, 12)

    # ATS Details
    create_score_card(
        doc,
        "ATS COMPLIANCE ANALYSIS",
        ats_score,
        analysis_results.get('ats_details', []),
        "2F5597"
    )

    # Match Details
    create_score_card(
        doc,
        "JOB MATCH ANALYSIS",
        match_score,
        analysis_results.get('match_details', []),
        "2F5597"
    )

    # Keywords Analysis
    add_heading_with_style(doc, "KEYWORDS ANALYSIS", 2, (23, 54, 93), 18, 12)

    keywords_table = doc.add_table(rows=1, cols=2)
    keywords_table.style = 'Table Grid'
    keywords_table.autofit = False
    keywords_table.columns[0].width = Inches(2.8)
    keywords_table.columns[1].width = Inches(3.2)

    # Header row
    header_cells = keywords_table.rows[0].cells
    header_cells[0].text = "Keywords Found"
    header_cells[1].text = "Keywords Missing"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "2F5597")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add keywords
    found_keywords = analysis_results.get('keywords_found', [])
    missing_keywords = analysis_results.get('keywords_missing', [])

    max_rows = max(len(found_keywords), len(missing_keywords), 1)

    for i in range(max_rows):
        if i >= len(keywords_table.rows) - 1:
            keywords_table.add_row()

        row_cells = keywords_table.rows[i + 1].cells

        if i < len(found_keywords):
            row_cells[0].text = found_keywords[i]
            set_cell_background(row_cells[0], "E8F5E9")
        else:
            row_cells[0].text = ""

        if i < len(missing_keywords):
            row_cells[1].text = missing_keywords[i]
            set_cell_background(row_cells[1], "FFEBEE")
        else:
            row_cells[1].text = ""

    doc.add_paragraph()

    # ---------------------- STRENGTHS & WEAKNESSES ----------------------
    add_heading_with_style(doc, "STRENGTHS & DEVELOPMENT AREAS", 1, (23, 54, 93), 24, 12)

    sw_table = doc.add_table(rows=1, cols=2)
    sw_table.style = 'Table Grid'
    sw_table.autofit = False
    sw_table.columns[0].width = Inches(2.8)
    sw_table.columns[1].width = Inches(3.2)

    # Header row
    sw_header_cells = sw_table.rows[0].cells
    sw_header_cells[0].text = "Key Strengths"
    sw_header_cells[1].text = "Areas for Development"

    for cell in sw_header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "2F5597")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add strengths and weaknesses
    strengths = analysis_results.get('strengths', [])
    weaknesses = analysis_results.get('weaknesses', [])

    max_sw_rows = max(len(strengths), len(weaknesses), 1)

    for i in range(max_sw_rows):
        if i >= len(sw_table.rows) - 1:
            sw_table.add_row()

        row_cells = sw_table.rows[i + 1].cells

        if i < len(strengths):
            row_cells[0].text = f"• {strengths[i]}"
            set_cell_background(row_cells[0], "E8F5E9")
        else:
            row_cells[0].text = ""

        if i < len(weaknesses):
            row_cells[1].text = f"• {weaknesses[i]}"
            set_cell_background(row_cells[1], "FFEBEE")
        else:
            row_cells[1].text = ""

    doc.add_paragraph()

    # ---------------------- IMPROVEMENT RECOMMENDATIONS ----------------------
    add_heading_with_style(doc, "IMPROVEMENT RECOMMENDATIONS", 1, (23, 54, 93), 24, 12)

    tips = analysis_results.get('improvement_tips', [])
    for i, tip in enumerate(tips, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(tip)
        run.font.size = Pt(11)

    # ---------------------- FINAL ASSESSMENT ----------------------
    add_heading_with_style(doc, "FINAL ASSESSMENT", 1, (23, 54, 93), 24, 12)

    final_note = doc.add_paragraph()
    final_note.paragraph_format.space_before = Pt(12)
    final_note.paragraph_format.space_after = Pt(18)
    final_note.paragraph_format.line_spacing = 1.15

    run = final_note.add_run("Based on our comprehensive analysis, this resume is ")
    run.font.size = Pt(11)

    if ats_score >= 80 and match_score >= 80:
        run = final_note.add_run("highly competitive ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 125, 50)
        final_note.add_run("for the target position. Minor refinements may further enhance its effectiveness.")
    elif ats_score >= 60 and match_score >= 60:
        run = final_note.add_run("moderately aligned ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(239, 108, 0)
        final_note.add_run("with the target position. Implementing the recommended changes will significantly improve candidacy.")
    else:
        run = final_note.add_run("not optimally prepared ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(198, 40, 40)
        final_note.add_run("for the target position. Substantial revisions are recommended before application.")

    # Add footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Confidential Report - Generated by PrepVista on {datetime.now().strftime('%Y-%m-%d')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Save document as DOCX
    try:
        doc.save(output_path)
        print(f"Report saved successfully: {output_path}")
        return output_path
    except Exception as e:
        print(f"Error saving document: {e}")
        return None