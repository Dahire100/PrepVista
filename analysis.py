# Install dependencies in terminal:
# pip install python-docx mediapipe opencv-python-headless transformers moviepy librosa matplotlib

import cv2
import mediapipe as mp
import numpy as np
import os
os.environ['CUDA_VISIBLE_DEVICES'] = '-1'
import librosa
from moviepy.editor import VideoFileClip
from transformers import pipeline
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
import time
from datetime import datetime
import tempfile
from contextlib import contextmanager

# ==============================================================================
# 1. ENHANCED PROFESSIONAL REPORT GENERATION (with original function name)
# ==============================================================================
def generate_fast_report(video_results, audio_results, filename="Fast_Analysis_Report.docx"):
    """Generate a comprehensive professional report - enhanced version"""
    doc = docx.Document()
    
    # Document styling
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Calculate overall score with weighted components
    score_components = {
        'posture': video_results['posture_score'] * 0.25,
        'eye_contact': video_results['eye_contact_score'] * 0.25,
        'confidence': video_results['confidence_score'] * 0.20,
        'gestures': (100 if video_results.get('gestures_ok', True) else 70) * 0.15,
        'vocal_clarity': (100 if audio_results and audio_results.get('filler_count', 0) < 5 else 70) * 0.15
    }
    
    overall_score = min(95, int(sum(score_components.values())))

    # Title Section
    title = doc.add_heading("PROFESSIONAL COMMUNICATION ANALYSIS REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header information
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Overall Score
    score_p = doc.add_paragraph()
    score_text = f"OVERALL SCORE: {overall_score}/100"
    score_run = score_p.add_run(score_text)
    score_run.bold = True
    score_run.font.size = Pt(14)
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Score interpretation
    if overall_score >= 85:
        interpretation = "Excellent - Professional level communication skills"
    elif overall_score >= 70:
        interpretation = "Good - Strong foundation with some areas for improvement"
    else:
        interpretation = "Developing - Focused practice needed in key areas"
    
    interpret_p = doc.add_paragraph()
    interpret_run = interpret_p.add_run(interpretation)
    interpret_run.italic = True
    interpret_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    summary = generate_executive_summary(video_results, audio_results, overall_score)
    doc.add_paragraph(summary)

    # Key Insights
    doc.add_heading("Key Insights", 1)
    insights = generate_key_insights(video_results, audio_results)
    for insight in insights:
        doc.add_paragraph(insight)

    # Performance Metrics
    doc.add_heading("Performance Metrics", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Metric"
    hdr_cells[2].text = "Score"

    metrics = [
        ("Body Language", "Posture", f"{video_results['posture_score']:.0f}%"),
        ("Body Language", "Eye Contact", f"{video_results['eye_contact_score']:.0f}%"),
        ("Body Language", "Confidence", f"{video_results['confidence_score']:.0f}%"),
    ]

    if audio_results:
        metrics.extend([
            ("Vocal Delivery", "Speaking Pace", f"{audio_results['wpm']:.0f} WPM"),
            ("Vocal Delivery", "Filler Words", f"{audio_results['filler_count']}"),
            ("Vocal Delivery", "Vocal Variety", audio_results['monotony'].split(' ')[0]),
        ])

    for category, metric, score in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = metric
        row_cells[2].text = score

    # Action Plan
    doc.add_heading("Action Plan", 1)
    actions = generate_action_plan(video_results, audio_results)
    for i, action in enumerate(actions[:5], 1):
        doc.add_paragraph(f"{i}. {action}")

    # Performance Visualization
    doc.add_heading("Performance Summary", 1)
    fig = create_performance_chart(video_results, audio_results)
    memfile = io.BytesIO()
    plt.savefig(memfile, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    doc.add_picture(memfile, width=Inches(6.5))

    doc.save(filename)
    print(f"✅ Report generated: {filename}")
    return overall_score

def generate_executive_summary(video_results, audio_results, overall_score):
    """Generate an executive summary"""
    summary_parts = []
    
    if overall_score >= 85:
        summary_parts.append("Your communication skills demonstrate strong professional competence.")
    elif overall_score >= 70:
        summary_parts.append("You have a solid foundation in communication with specific areas for growth.")
    else:
        summary_parts.append("Focus on fundamental communication skills for maximum impact.")
    
    # Video strengths
    strong_areas = []
    if video_results['posture_score'] >= 75:
        strong_areas.append("confident posture")
    if video_results['eye_contact_score'] >= 70:
        strong_areas.append("strong audience engagement")
    
    if strong_areas:
        summary_parts.append(f"Key strengths include {', '.join(strong_areas)}.")
    
    # Development areas
    development_areas = []
    if video_results['posture_score'] < 65:
        development_areas.append("posture alignment")
    if video_results['eye_contact_score'] < 60:
        development_areas.append("consistent eye contact")
    if audio_results and audio_results.get('filler_count', 0) > 5:
        development_areas.append("reducing filler words")
    
    if development_areas:
        summary_parts.append(f"Focus on improving {', '.join(development_areas)} for enhanced effectiveness.")
    
    return " ".join(summary_parts)

def generate_key_insights(video_results, audio_results):
    """Generate key insights based on analysis results"""
    insights = []
    
    # Video insights
    if video_results['posture_score'] > 75:
        insights.append("✅ Strong, confident posture with good alignment")
    elif video_results['posture_score'] > 55:
        insights.append("⚡ Good posture foundation - minor adjustments needed")
    else:
        insights.append("📉 Work on maintaining upright posture for better presence")

    if video_results['eye_contact_score'] > 70:
        insights.append("✅ Excellent audience engagement through eye contact")
    elif video_results['eye_contact_score'] > 50:
        insights.append("⚡ Decent eye contact - aim for more consistency")
    else:
        insights.append("📉 Improve eye contact for better audience connection")

    if video_results['confidence_score'] > 70:
        insights.append("✅ Confident and professional presence")
    else:
        insights.append("📉 Build confidence through practice and preparation")

    # Audio insights
    if audio_results:
        if 120 <= audio_results['wpm'] <= 150:
            insights.append("✅ Ideal speaking pace for clear communication")
        elif audio_results['wpm'] > 150:
            insights.append("📉 Speaking too fast - slow down for better clarity")
        else:
            insights.append("📉 Speaking too slow - increase energy and pace")

        if audio_results['filler_count'] < 5:
            insights.append("✅ Clear speech with minimal filler words")
        else:
            insights.append("📉 Reduce filler words for more polished delivery")

        if audio_results['monotony'] != "High":
            insights.append("✅ Good vocal variety and expression")
        else:
            insights.append("📉 Add more vocal inflection to maintain interest")

    return insights

def generate_action_plan(video_results, audio_results):
    """Generate personalized action plan"""
    actions = []
    
    # Posture actions
    if video_results['posture_score'] < 70:
        actions.append("Practice standing straight against a wall daily for 5 minutes")
    else:
        actions.append("Maintain excellent posture through regular awareness checks")
    
    # Eye contact actions
    if video_results['eye_contact_score'] < 65:
        actions.append("Record short videos practicing direct eye contact with the camera")
    else:
        actions.append("Continue strong eye contact practice in various scenarios")
    
    # Confidence actions
    if video_results['confidence_score'] < 70:
        actions.append("Practice power poses before important speaking engagements")
    
    # Audio actions
    if audio_results:
        if audio_results['wpm'] > 160:
            actions.append("Use a metronome app to practice slower, more deliberate pacing")
        elif audio_results['wpm'] < 110:
            actions.append("Practice reading aloud with energy to increase speaking pace")
        
        if audio_results['filler_count'] > 5:
            actions.append("Practice speaking with deliberate pauses instead of filler words")
    
    # General best practices
    actions.extend([
        "Practice 2-minute presentations daily",
        "Watch recordings to self-assess and identify improvement areas",
        "Get feedback from colleagues or mentors regularly"
    ])
    
    return actions[:5]

def create_performance_chart(video_results, audio_results):
    """Create performance visualization"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 4))
    
    # Main metrics
    categories = ['Posture', 'Eye Contact', 'Confidence']
    scores = [video_results['posture_score'], video_results['eye_contact_score'], video_results['confidence_score']]
    colors = ['#4A90E2', '#50E3C2', '#F5A623']
    
    bars = ax1.bar(categories, scores, color=colors, alpha=0.8)
    ax1.set_ylim(0, 100)
    ax1.set_ylabel('Score (%)', fontweight='bold')
    ax1.set_title('Core Communication Metrics', fontweight='bold')
    for bar in bars:
        height = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2., height + 1, f'{height:.0f}%', 
                ha='center', va='bottom', fontweight='bold')
    
    # Audio metrics if available
    if audio_results:
        audio_categories = ['Pace', 'Clarity', 'Variety']
        audio_scores = [
            min(100, max(0, 160 - audio_results['wpm']) * 2),  # Pace score
            max(0, 100 - (audio_results['filler_count'] * 10)),  # Clarity score
            {'High': 30, 'Medium': 70, 'Low': 90}[audio_results['monotony']]  # Variety score
        ]
        ax2.bar(audio_categories, audio_scores, color=['#9B59B6', '#E74C3C', '#3498DB'], alpha=0.8)
        ax2.set_ylim(0, 100)
        ax2.set_title('Vocal Delivery Metrics', fontweight='bold')
        ax2.set_ylabel('Performance Score', fontweight='bold')
    
    plt.tight_layout()
    return fig

# ==============================================================================
# 2. OPTIMIZED ANALYSIS FUNCTIONS (with original names)
# ==============================================================================
@contextmanager
def mediapipe_context():
    """Context manager for MediaPipe resources"""
    mp_pose = mp.solutions.pose.Pose(
        min_detection_confidence=0.5, 
        min_tracking_confidence=0.5, 
        model_complexity=0
    )
    mp_face_mesh = mp.solutions.face_mesh.FaceMesh(
        max_num_faces=1, 
        min_detection_confidence=0.5, 
        min_tracking_confidence=0.5,
        static_image_mode=False
    )
    try:
        yield mp_pose, mp_face_mesh
    finally:
        mp_pose.close()
        mp_face_mesh.close()

def analyze_audio_ultrafast(audio_path):
    """Optimized audio analysis - enhanced version"""
    print("--- 🎤 Ultra-fast Audio Analysis ---")
    start_time = time.time()
    
    try:
        # Load audio with optimized parameters
        y, sr = librosa.load(audio_path, sr=8000, duration=60)
        duration = librosa.get_duration(y=y, sr=sr)
        
        if duration < 3:
            return None
        
        # Use whisper for speech recognition
        whisper_pipe = pipeline("automatic-speech-recognition", model="openai/whisper-tiny", device=-1)
        max_duration = min(30, duration)
        if duration > 30:
            y = y[:int(30 * sr)]
            
        result = whisper_pipe(y)
        text = result["text"].strip()
        
        if not text:
            return None
            
        word_count = len(text.split())
        wpm = (word_count / max_duration) * 60
        
        # Filler word detection
        filler_words = [' um ', ' uh ', ' like ', ' you know ', ' actually ', ' basically ']
        filler_count = sum(text.lower().count(filler) for filler in filler_words)
        
        # Pitch analysis
        pitches = librosa.pyin(y, fmin=50, fmax=400, sr=sr, frame_length=512)[0]
        pitches = pitches[~np.isnan(pitches)]
        pitch_variation = np.std(pitches) if len(pitches) > 10 else 0
        monotony = "High" if pitch_variation < 30 else "Medium" if pitch_variation < 60 else "Low"
        
        print(f"✅ Audio analysis completed in {time.time()-start_time:.1f}s")
        return {
            "wpm": wpm, 
            "filler_count": filler_count, 
            "monotony": monotony, 
            "duration": duration
        }
        
    except Exception as e:
        print(f"⚠ Audio analysis skipped: {e}")
        return None

def analyze_video_ultrafast(video_path):
    """Optimized video analysis - enhanced version"""
    print("--- 🎥 Ultra-fast Video Analysis ---")
    start_time = time.time()
    
    with mediapipe_context() as (pose, face_mesh):
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            return {}

        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        
        # Adaptive sampling
        target_samples = min(60, total_frames)
        frame_interval = max(1, total_frames // target_samples)

        posture_samples = []
        eye_contact_samples = []
        confidence_samples = []
        frame_count = 0

        for i in range(0, total_frames, frame_interval):
            if frame_count >= target_samples:
                break
                
            cap.set(cv2.CAP_PROP_POS_FRAMES, i)
            ret, frame = cap.read()
            if not ret:
                break
                
            frame_count += 1
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            
            # Pose analysis
            pose_results = pose.process(rgb_frame)
            if pose_results.pose_landmarks:
                landmarks = pose_results.pose_landmarks.landmark
                
                # Posture: shoulder alignment
                left_shoulder = landmarks[mp.solutions.pose.PoseLandmark.LEFT_SHOULDER.value]
                right_shoulder = landmarks[mp.solutions.pose.PoseLandmark.RIGHT_SHOULDER.value]
                shoulder_diff = abs(left_shoulder.y - right_shoulder.y)
                posture_samples.append(shoulder_diff < 0.08)
                
                # Confidence: open body language
                left_hip = landmarks[mp.solutions.pose.PoseLandmark.LEFT_HIP.value]
                right_hip = landmarks[mp.solutions.pose.PoseLandmark.RIGHT_HIP.value]
                shoulder_width = abs(left_shoulder.x - right_shoulder.x)
                hip_width = abs(left_hip.x - right_hip.x)
                confidence_samples.append(shoulder_width > hip_width * 0.8)
            
            # Face analysis every 3rd frame for performance
            if frame_count % 3 == 0:
                face_results = face_mesh.process(rgb_frame)
                if face_results.multi_face_landmarks:
                    landmarks = face_results.multi_face_landmarks[0].landmark
                    left_eye = landmarks[33]
                    right_eye = landmarks[263]
                    eye_distance = abs(left_eye.x - right_eye.x)
                    eye_contact_samples.append(eye_distance > 0.1)

        cap.release()

        # Calculate scores with smoothing
        posture_score = calculate_smoothed_score(posture_samples, 65, 5)
        eye_contact_score = calculate_smoothed_score(eye_contact_samples, 60, 5)
        confidence_score = calculate_smoothed_score(confidence_samples, 65, 5)

        print(f"✅ Video analysis completed in {time.time()-start_time:.1f}s")
        return {
            "posture_score": posture_score,
            "eye_contact_score": eye_contact_score, 
            "confidence_score": confidence_score,
            "blinks_ok": True, 
            "gestures_ok": True, 
            "frames_analyzed": frame_count
        }

def calculate_smoothed_score(samples, baseline, variation):
    """Calculate score with baseline smoothing"""
    if not samples:
        return baseline
    
    raw_score = (sum(samples) / len(samples)) * 100
    # Apply mild smoothing to avoid extreme scores
    return min(95, max(30, raw_score))

# ==============================================================================
# 3. MAIN EXECUTION FUNCTIONS
# ==============================================================================
@contextmanager
def temporary_audio_file(video_path):
    """Context manager for temporary audio file handling"""
    temp_audio = None
    try:
        if VideoFileClip(video_path).audio:
            temp_audio_file = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
            temp_audio = temp_audio_file.name
            temp_audio_file.close()
            
            # Fast audio extraction
            with VideoFileClip(video_path) as clip:
                clip.audio.write_audiofile(
                    temp_audio, 
                    verbose=False, 
                    logger=None
                )
            yield temp_audio
        else:
            yield None
    finally:
        # Cleanup with retry logic
        if temp_audio and os.path.exists(temp_audio):
            for attempt in range(3):
                try:
                    time.sleep(0.1 * attempt)
                    os.remove(temp_audio)
                    break
                except (PermissionError, OSError):
                    if attempt == 2:
                        print(f"⚠️ Could not delete temporary audio file: {temp_audio}")

def run_ultrafast_analysis(video_path):
    """Main analysis function with optimized performance"""
    if not os.path.exists(video_path):
        print(f"❌ Video file not found: {video_path}")
        return
    
    total_start = time.time()
    print("🚀 STARTING ULTRA-FAST COMMUNICATION ANALYSIS")
    print("=" * 50)
    
    # Audio analysis
    audio_results = None
    with temporary_audio_file(video_path) as temp_audio:
        if temp_audio:
            audio_results = analyze_audio_ultrafast(temp_audio)
        else:
            print("⚠️ No audio track found in video")
    
    # Video analysis
    video_results = analyze_video_ultrafast(video_path)
    
    # Generate enhanced report
    overall_score = generate_fast_report(video_results, audio_results, "Enhanced_Analysis_Report.docx")
    
    # Performance summary
    total_time = time.time() - total_start
    print("=" * 50)
    print(f"🎉 ANALYSIS COMPLETE IN {total_time:.1f} SECONDS!")
    print(f"📊 Overall Score: {overall_score}/100")
    print(f"💾 Enhanced Report: Enhanced_Analysis_Report.docx")
    print("=" * 50)
    
    return overall_score

# ==============================================================================
# 4. QUICK ANALYSIS MODE
# ==============================================================================
def one_click_analysis(video_path):
    """Fast analysis for quick feedback"""
    if not os.path.exists(video_path):
        print("❌ Video file not found")
        return
    
    start_time = time.time()
    print("⚡ ONE-CLICK INSTANT ANALYSIS MODE - Processing...")
    
    with mediapipe_context() as (pose, face_mesh):
        cap = cv2.VideoCapture(video_path)
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        sample_frames = min(20, total_frames)
        interval = max(1, total_frames // sample_frames)
        
        posture_count = 0
        frontal_face_count = 0
        frames_checked = 0
        
        for i in range(0, total_frames, interval):
            if frames_checked >= sample_frames:
                break
                
            cap.set(cv2.CAP_PROP_POS_FRAMES, i)
            ret, frame = cap.read()
            if not ret:
                break
                
            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            
            # Quick pose check
            pose_result = pose.process(rgb)
            if pose_result.pose_landmarks:
                landmarks = pose_result.pose_landmarks.landmark
                if abs(landmarks[11].y - landmarks[12].y) < 0.08:  # Shoulder alignment
                    posture_count += 1
            
            # Quick face check
            face_result = face_mesh.process(rgb)
            if face_result.multi_face_landmarks:
                frontal_face_count += 1
                
            frames_checked += 1
        
        cap.release()
    
    # Calculate quick scores
    posture_score = (posture_count / frames_checked) * 100
    eye_contact_score = (frontal_face_count / frames_checked) * 100
    confidence_score = min(100, (posture_score + eye_contact_score) / 2 + 15)
    
    video_results = {
        "posture_score": posture_score, 
        "eye_contact_score": eye_contact_score, 
        "confidence_score": confidence_score, 
        "blinks_ok": True, 
        "gestures_ok": True
    }
    
    # Generate quick report
    generate_fast_report(video_results, None, "Instant_Analysis_Report.docx")
    print(f"🎉 INSTANT ANALYSIS COMPLETE IN {time.time()-start_time:.1f} SECONDS!")
    return video_results

# ==============================================================================
# 5. MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    video_path = "video.mp4"
    
    if not os.path.exists(video_path):
        print(f"❌ Video file not found: {video_path}")
    else:
        print("Select analysis mode:")
        print("1. Ultra-Fast Analysis (30-60 seconds) - Enhanced comprehensive report")
        print("2. One-Click Instant Analysis (10-20 seconds) - Fast feedback")
        
        choice = input("Enter choice (1 or 2): ").strip()
        
        if choice == "2":
            one_click_analysis(video_path)
        else:
            run_ultrafast_analysis(video_path)